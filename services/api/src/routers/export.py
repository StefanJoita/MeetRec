# services/api/src/routers/export.py
# ============================================================
# Export Router — PDF / DOCX / TXT
# ============================================================

import io
import re
import uuid
from datetime import datetime
from urllib.parse import quote

from fastapi import APIRouter, Depends, HTTPException, Query, Request
from fastapi.responses import StreamingResponse
from slowapi import Limiter
from slowapi.util import get_remote_address
from sqlalchemy.ext.asyncio import AsyncSession

from src.database import get_db
from src.middleware.audit import log_audit
from src.middleware.auth import get_current_user, check_recording_access
from src.models.audit_log import User
from src.models.recording import Recording
from src.services.transcript_service import TranscriptService
from sqlalchemy import select

router = APIRouter(prefix="/export", tags=["export"])
limiter = Limiter(key_func=get_remote_address)


def _safe_filename(title: str) -> str:
    """Sanitizează titlul pentru a fi folosit sigur în header-ul Content-Disposition.
    Elimină orice caracter care ar putea cauza header injection (CRLF, ghilimele, etc.).
    """
    # Eliminăm caracterele de control (inclusiv CR/LF) și ghilimelele
    safe = re.sub(r'[\x00-\x1f\x7f"\\]', '', title)
    # Înlocuim spațiile cu underscore
    safe = safe.strip().replace(' ', '_')
    # Păstrăm maxim 50 de caractere
    safe = safe[:50]
    return safe or 'export'


def _format_time(seconds: float) -> str:
    """Convertește secunde în format MM:SS."""
    total = int(seconds)
    m, s = divmod(total, 60)
    h, m = divmod(m, 60)
    if h:
        return f"{h:02d}:{m:02d}:{s:02d}"
    return f"{m:02d}:{s:02d}"


async def _build_speaker_display_map(
    segments, speaker_mapping: dict, db: AsyncSession
) -> dict[str, str]:
    """
    Construiește un dict {speaker_id → nume afișabil} pentru toate segmentele.

    Rezolvă două cazuri:
      1. speaker_id este deja un UUID de user (post-rezolvare directă)
      2. speaker_id este "SPEAKER_XX" → speaker_mapping["SPEAKER_XX"] = user_uuid → nume user
    """
    speaker_ids = {seg.speaker_id for seg in segments if seg.speaker_id}
    if not speaker_ids:
        return {}

    display_map: dict[str, str] = {}
    # speaker_id (cheie în segmente) → UUID de user de rezolvat
    to_resolve: dict[str, uuid.UUID] = {}

    for sid in speaker_ids:
        try:
            # Cazul 1: speaker_id este direct un UUID de user
            to_resolve[sid] = uuid.UUID(sid)
        except ValueError:
            # Cazul 2: "SPEAKER_XX" — căutăm în speaker_mapping
            raw = speaker_mapping.get(sid)
            if raw:
                try:
                    to_resolve[sid] = uuid.UUID(raw)
                except ValueError:
                    display_map[sid] = raw  # valoare neparsabilă, o afișăm direct

    if to_resolve:
        unique_uuids = list({str(u) for u in to_resolve.values()})
        result = await db.execute(
            select(User).where(User.id.in_([uuid.UUID(u) for u in unique_uuids]))
        )
        user_names = {str(u.id): (u.full_name or u.username) for u in result.scalars().all()}
        for sid, uid in to_resolve.items():
            name = user_names.get(str(uid))
            if name:
                display_map[sid] = name

    return display_map


async def _get_transcript_and_recording(
    recording_id: uuid.UUID,
    db: AsyncSession,
):
    """Obține transcriptul și înregistrarea. Ridică 404 dacă nu există."""
    rec_result = await db.execute(
        select(Recording).where(Recording.id == recording_id)
    )
    recording = rec_result.scalar_one_or_none()
    if not recording:
        raise HTTPException(status_code=404, detail="Înregistrarea nu există.")

    service = TranscriptService(db)
    transcript = await service.get_by_recording_id(recording_id)
    if not transcript or transcript.status != "completed":
        raise HTTPException(
            status_code=404,
            detail="Transcriptul nu este disponibil sau nu a fost finalizat.",
        )
    return recording, transcript


@router.get(
    "/recording/{recording_id}",
    summary="Exportă transcriptul (PDF / DOCX / TXT)",
)
@limiter.limit("20/hour")
async def export_transcript(
    request: Request,
    recording_id: uuid.UUID,
    format: str = Query(default="txt", pattern="^(pdf|docx|txt)$"),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """
    Exportă transcriptul unei înregistrări în formatul dorit.

    - **pdf** — document PDF formatat, cu titlu și timestamps
    - **docx** — document Word cu stiluri
    - **txt** — text simplu, ușor de prelucrat
    """
    recording, transcript = await _get_transcript_and_recording(recording_id, db)

    if not await check_recording_access(recording_id, current_user, db):
        raise HTTPException(status_code=403, detail="Acces interzis la această înregistrare.")

    await log_audit(
        request, db,
        action="EXPORT",
        resource_type="recording",
        resource_id=recording_id,
        details={"format": format},
    )

    filename_base = _safe_filename(recording.title)
    display_map = await _build_speaker_display_map(
        transcript.segments, recording.speaker_mapping or {}, db
    )

    if format == "txt":
        return _export_txt(recording, transcript, filename_base, display_map)
    elif format == "pdf":
        return _export_pdf(recording, transcript, filename_base, display_map)
    elif format == "docx":
        return _export_docx(recording, transcript, filename_base, display_map)


# ── TXT ──────────────────────────────────────────────────────

def _export_txt(recording, transcript, filename_base: str, display_map: dict) -> StreamingResponse:
    lines = [
        f"TRANSCRIERE: {recording.title}",
        f"Data ședinței: {recording.meeting_date}",
        f"Durată: {recording.duration_formatted}",
        f"Limbă: {transcript.language or 'necunoscută'}",
        f"Exportat: {datetime.now().strftime('%d.%m.%Y %H:%M')}",
        "",
        "=" * 60,
        "",
    ]
    for seg in transcript.segments:
        timestamp = f"[{_format_time(float(seg.start_time))}]"
        label = display_map.get(seg.speaker_id) if seg.speaker_id else None
        if label:
            lines.append(f"{timestamp} {label}: {seg.text}")
        else:
            lines.append(f"{timestamp}  {seg.text}")

    content = "\n".join(lines)
    buf = io.BytesIO(content.encode("utf-8"))

    encoded = quote(f"{filename_base}.txt")
    return StreamingResponse(
        buf,
        media_type="text/plain; charset=utf-8",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded}"},
    )


# ── PDF ──────────────────────────────────────────────────────

def _register_pdf_fonts() -> tuple[str, str]:
    """
    Înregistrează fonturile DejaVu (suport Unicode complet, inclusiv diacritice românești).
    Returnează (font_normal, font_bold).
    Fallback la Helvetica dacă fonturile nu sunt disponibile în container.
    """
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    font_path_normal = "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"
    font_path_bold = "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"

    import os
    if os.path.exists(font_path_normal) and os.path.exists(font_path_bold):
        pdfmetrics.registerFont(TTFont("DejaVuSans", font_path_normal))
        pdfmetrics.registerFont(TTFont("DejaVuSans-Bold", font_path_bold))
        return "DejaVuSans", "DejaVuSans-Bold"
    return "Helvetica", "Helvetica-Bold"


def _export_pdf(recording, transcript, filename_base: str, display_map: dict) -> StreamingResponse:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle

    font_normal, font_bold = _register_pdf_fonts()

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=2.5 * cm,
        rightMargin=2.5 * cm,
        topMargin=2.5 * cm,
        bottomMargin=2.5 * cm,
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "CustomTitle",
        parent=styles["Title"],
        fontSize=16,
        spaceAfter=6,
        fontName=font_bold,
    )
    meta_style = ParagraphStyle(
        "Meta",
        parent=styles["Normal"],
        fontSize=9,
        textColor=colors.gray,
        spaceAfter=2,
        fontName=font_normal,
    )
    segment_style = ParagraphStyle(
        "Segment",
        parent=styles["Normal"],
        fontSize=10,
        spaceAfter=4,
        leading=14,
        fontName=font_normal,
    )
    timestamp_style = ParagraphStyle(
        "Timestamp",
        parent=styles["Normal"],
        fontSize=8,
        textColor=colors.HexColor("#666666"),
        spaceAfter=1,
        fontName=font_normal,
    )
    speaker_style = ParagraphStyle(
        "Speaker",
        parent=styles["Normal"],
        fontSize=8,
        textColor=colors.HexColor("#1a56a0"),
        spaceAfter=1,
        fontName=font_bold,
    )

    story = []

    # Titlu
    story.append(Paragraph(recording.title, title_style))
    story.append(Paragraph(f"Data ședinței: {recording.meeting_date}", meta_style))
    story.append(Paragraph(f"Durată: {recording.duration_formatted}", meta_style))
    story.append(Paragraph(
        f"Limbă: {transcript.language or 'necunoscută'} · "
        f"Cuvinte: {transcript.word_count} · "
        f"Exportat: {datetime.now().strftime('%d.%m.%Y %H:%M')}",
        meta_style,
    ))
    story.append(Spacer(1, 0.5 * cm))

    # Linie separatoare
    story.append(Table(
        [[""]],
        colWidths=["100%"],
        style=TableStyle([
            ("LINEBELOW", (0, 0), (-1, 0), 0.5, colors.lightgrey),
        ]),
    ))
    story.append(Spacer(1, 0.4 * cm))

    # Segmente
    for seg in transcript.segments:
        ts = _format_time(float(seg.start_time))
        label = display_map.get(seg.speaker_id) if seg.speaker_id else None
        if label:
            story.append(Paragraph(f"{ts} · {label}", speaker_style))
        else:
            story.append(Paragraph(ts, timestamp_style))
        story.append(Paragraph(seg.text, segment_style))

    doc.build(story)
    buf.seek(0)

    encoded = quote(f"{filename_base}.pdf")
    return StreamingResponse(
        buf,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded}"},
    )


# ── DOCX ─────────────────────────────────────────────────────

def _export_docx(recording, transcript, filename_base: str, display_map: dict) -> StreamingResponse:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Margini
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Titlu
    title_par = doc.add_heading(recording.title, level=1)
    title_par.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Metadata
    meta = doc.add_paragraph()
    meta.add_run(f"Data ședinței: {recording.meeting_date}\n").font.size = Pt(9)
    meta.add_run(f"Durată: {recording.duration_formatted}\n").font.size = Pt(9)
    run = meta.add_run(
        f"Limbă: {transcript.language or 'necunoscută'} · "
        f"Cuvinte: {transcript.word_count} · "
        f"Exportat: {datetime.now().strftime('%d.%m.%Y %H:%M')}"
    )
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.add_paragraph()  # spațiu

    # Segmente
    for seg in transcript.segments:
        ts = _format_time(float(seg.start_time))
        label = display_map.get(seg.speaker_id) if seg.speaker_id else None

        ts_par = doc.add_paragraph()
        if label:
            ts_run = ts_par.add_run(f"[{ts}] {label}")
            ts_run.font.color.rgb = RGBColor(0x1A, 0x56, 0xA0)
        else:
            ts_run = ts_par.add_run(f"[{ts}]")
            ts_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        ts_run.font.size = Pt(8)
        ts_run.bold = True

        text_par = doc.add_paragraph(seg.text)
        text_par.style.font.size = Pt(10)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    encoded = quote(f"{filename_base}.docx")
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded}"},
    )
